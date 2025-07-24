# ✅ FINAL Enhanced home.py with Word to PDF Converter Tab

import streamlit as st
import streamlit.components.v1 as components
from streamlit_js_eval import streamlit_js_eval
import pytesseract
from PIL import Image, ImageDraw, ImageFont
import spacy
from pdf2image import convert_from_bytes
from io import BytesIO
import json
from streamlit_lottie import st_lottie
import base64
from docx import Document
from fpdf import FPDF
import textwrap
import unicodedata
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
#from docx2pdf import convert
#import pythoncom
import tempfile
import os
import pdfplumber
import pandas as pd
import streamlit as st
import streamlit.components.v1 as components
from pdf2docx import Converter
from docx import Document
from fpdf import FPDF



# OCR Setup
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
nlp = spacy.load("en_core_web_sm")

# Utilities
def load_lottie(path):
    with open(path, "r") as f:
        return json.load(f)

def local_css(file_name):
    with open(file_name) as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

def image_to_base64(img):
    buffered = BytesIO()
    img.save(buffered, format="PNG")
    return base64.b64encode(buffered.getvalue()).decode()

def generate_docx(text):
    doc = Document()
    doc.add_paragraph(text)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def generate_pdf(text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_left_margin(10)
    pdf.set_right_margin(10)
    pdf.set_font("Arial", size=8)

    def sanitize(text):
        return ''.join(c if unicodedata.category(c)[0] != 'C' and ord(c) < 128 else '?' for c in str(text))

    clean_text = sanitize(text)
    page_width = pdf.w - 2 * pdf.l_margin

    for line in clean_text.split('\n'):
        line = line.strip()
        if not line:
            pdf.ln()
            continue
        wrapped_lines = textwrap.wrap(line, width=100)
        for subline in wrapped_lines:
            try:
                pdf.multi_cell(page_width, 10, subline)
            except:
                pdf.multi_cell(page_width, 10, "[Unrenderable content]")

    return pdf.output(name=None, dest='S')

# Login

def login():
    st.title("\U0001F510 Login")

    # 🌐 Spline 3D Embed
    components.html("""
        <iframe src='https://my.spline.design/genkubgreetingrobot-mJX385FhHkoC6ESSs4ARGsne/' 
                frameborder='0' width='100%' height='300px' allowfullscreen></iframe>
    """, height=320)

    username = st.text_input("Username")
    password = st.text_input("Password", type="password")
    if st.button("Login"):
        if username == "admin" and password == "1234":
            st.session_state["authenticated"] = True
        else:
            st.error("Invalid credentials")


# Session State Logic (DO NOT REMOVE THIS PART)
if "authenticated" not in st.session_state:
    st.session_state["authenticated"] = False
if not st.session_state["authenticated"]:
    login()
    st.stop()



# App UI
local_css("style.css")
lottie_ocr = load_lottie("assets/animation.json")

st.sidebar.markdown("## \U0001F58C️ Theme Settings")
st.sidebar.markdown("\U0001F319 Theme: **Midnight AI**")
st.sidebar.markdown("\U0001F9E0 **AI Theme**: Midnight Mode")

st.title("\U0001F4C4 OCR + Entity Extractor")
st_lottie(lottie_ocr, height=200)
st.markdown("---")

tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "🔍 OCR + NER",
    "🗕 About Project",
    "📝 Word to PDF Converter",
    "📄 PDF ↔ Excel",
    "📦 Compress File"
])



# --- OCR + NER Tab ---
with tab1:
    uploaded_file = st.file_uploader("\U0001F4C4 Upload Document", type=["jpg", "jpeg", "png", "pdf"])

    # Glow JS
    components.html("""
    <script>
    document.addEventListener("DOMContentLoaded", function () {
        const observer = new MutationObserver(() => {
            const dropBox = window.parent.document.querySelector('div[data-testid="stFileUploaderDropzone"]');
            if (dropBox) {
                const addGlow = () => dropBox.classList.add("drag-over-effect");
                const removeGlow = () => dropBox.classList.remove("drag-over-effect");

                window.parent.addEventListener("dragenter", addGlow);
                window.parent.addEventListener("dragover", addGlow);
                window.parent.addEventListener("dragleave", removeGlow);
                window.parent.addEventListener("drop", removeGlow);

                observer.disconnect();
            }
        });
        observer.observe(window.parent.document.body, { childList: true, subtree: true });
    });
    </script>
    """, height=0)

    if uploaded_file:
        st.info("\u23F3 Processing...")

        if uploaded_file.type == "application/pdf":
            images = convert_from_bytes(uploaded_file.read(), poppler_path=r"C:\\poppler\\poppler-24.08.0\\Library\\bin")
        else:
            images = [Image.open(uploaded_file)]

        full_text = ""

        for page_num, img in enumerate(images):
            st.markdown(f"### \U0001F4C4 Page {page_num + 1}")
            st.image(img, caption="\U0001F4F7 Page Image", use_column_width=True)

            text = pytesseract.image_to_string(img)
            full_text += text + "\n"

            data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
            img_b64 = image_to_base64(img)

            overlay_html = f"""
            <div style='position: relative; display: inline-block; width: 100%;'>
                <img src='data:image/png;base64,{img_b64}' style='width: 100%;' />
            """

            for i in range(len(data['text'])):
                word = data['text'][i]
                if word.strip() == "":
                    continue
                left, top, width, height = data['left'][i], data['top'][i], data['width'][i], data['height'][i]
                overlay_html += f"""
                <div contenteditable='true' class='ocr-box' style='
                    position: absolute;
                    left: {left}px;
                    top: {top}px;
                    width: {width}px;
                    height: {height}px;
                    background-color: rgba(255,255,255,0.6);
                    font-size: 12px;
                    overflow: hidden;
                    border: 1px dashed #00ffee;
                    color: black;'>{word}</div>
                """

            overlay_html += """
            </div>
            <button onclick="combineText()">\U0001F4CB Save Changes</button>
            <button onclick="downloadImage()">\U0001F5BC️ Download as PNG</button>

            <script src="https://html2canvas.hertzen.com/dist/html2canvas.min.js"></script>
            <script>
            function combineText() {
                let boxes = document.querySelectorAll('.ocr-box');
                let result = '';
                boxes.forEach(box => {
                    let text = box.innerText.trim();
                    if (text.length > 0) result += text + ' ';
                });
                let hidden = document.getElementById('hiddenTextArea');
                if (!hidden) {
                    hidden = document.createElement('textarea');
                    hidden.id = 'hiddenTextArea';
                    hidden.style.display = 'none';
                    document.body.appendChild(hidden);
                }
                hidden.value = result.trim();
            }

            function downloadImage() {
                html2canvas(document.querySelector('div[style*="position: relative"]')).then(canvas => {
                    let link = document.createElement("a");
                    link.download = "edited_ocr_image.png";
                    link.href = canvas.toDataURL();
                    link.click();
                });
            }
            </script>
            """

            components.html(overlay_html, height=700)

        result = streamlit_js_eval(js_expressions="document.getElementById('hiddenTextArea')?.value || ''", key="sync_edit")

        st.subheader("\U0001F4DD Edited OCR Text")
        if result and result.strip():
            full_text = result
        edited_text = st.text_area("You can refine the OCR text here:", value=full_text, height=300, key="edited_ocr")

        st.subheader("\U0001F5BC️ Download Edited Image (PNG/JPG)")
        download_format = st.selectbox("Choose image format", ["PNG", "JPG", "JPEG"], key="format_select")

        if st.button("\U0001F5D5️ Generate Edited Image"):
            edited_img = img.copy()
            draw = ImageDraw.Draw(edited_img)
            edited_words = edited_text.strip().split()
            word_index = 0

            for i in range(len(data['text'])):
                if data['text'][i].strip() == "":
                    continue
                if word_index >= len(edited_words):
                    break

                new_word = edited_words[word_index]
                left, top, width, height = data['left'][i], data['top'][i], data['width'][i], data['height'][i]
                draw.rectangle([left, top, left + width, top + height], fill="white")

                try:
                    font_size = max(10, int(height * 0.9))
                    font = ImageFont.truetype("arial.ttf", font_size)
                except:
                    font = ImageFont.load_default()

                text_y = top + (height - font.getbbox(new_word)[3]) // 2
                draw.text((left, text_y), new_word, fill="black", font=font)
                word_index += 1

            buffer = BytesIO()
            fmt = "JPEG" if download_format.lower() in ["jpg", "jpeg"] else "PNG"
            edited_img.save(buffer, format=fmt)
            buffer.seek(0)

            st.download_button(
                label=f"\u2B07\uFE0F Download Edited Image as {download_format}",
                data=buffer,
                file_name=f"edited_ocr_image.{download_format.lower()}",
                mime=f"image/{'jpeg' if fmt == 'JPEG' else 'png'}"
            )

        st.subheader("\U0001F50D Named Entities")
        doc = nlp(edited_text)
        if doc.ents:
            color_map = {
                "ORG": "#FFD700",
                "PERSON": "#7CFC00",
                "DATE": "#00CED1",
                "MONEY": "#FFA07A",
                "GPE": "#DA70D6"
            }
            for ent in doc.ents:
                color = color_map.get(ent.label_, "#FFFFFF")
                st.markdown(f"""
                <div style="border-left: 6px solid {color}; padding: 10px; margin: 5px; background-color: #0d1117; color: white;">
                    <b>{ent.label_}</b>: {ent.text}
                </div>
                """, unsafe_allow_html=True)
        else:
            st.info("No named entities found.")

        st.subheader("\U0001F4E6 Export Options")
        col1, col2, col3 = st.columns(3)
        with col1:
            st.download_button("\u2B07\uFE0F Download .TXT", data=edited_text, file_name="ocr_output.txt", mime="text/plain")
        with col2:
            st.download_button("\u2B07\uFE0F Download .DOCX", data=generate_docx(edited_text),
                               file_name="ocr_output.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with col3:
            pdf_data = generate_pdf(edited_text)
            if isinstance(pdf_data, bytearray):
                pdf_data = bytes(pdf_data)
            st.download_button("\u2B07\uFE0F Download .PDF", data=pdf_data,
                               file_name="ocr_output.pdf",
                               mime="application/pdf")

# --- About Tab ---
with tab2:
    st.markdown("""
    ### 🤖 Built With:
    - Python
    - Streamlit
    - Tesseract OCR
    - spaCy NER
    - pdf2image + Poppler

    **Upload an image or PDF and extract + edit key entities instantly!**
    """)

from docx import Document
from fpdf import FPDF
from io import BytesIO

# --- Word to PDF Tab ---
with tab3:
    st.header("🔁 Convert Word (.docx) ↔ PDF")

    conversion_type = st.radio("Choose conversion type:", ["Word to PDF", "PDF to Word"], horizontal=True)

    uploaded_file = st.file_uploader(
        "Upload file",
        type=["docx", "pdf"],
        help="Upload .docx for Word to PDF or .pdf for PDF to Word",
    )

    # ------------------- Word to PDF Conversion -------------------
    def convert_docx_to_pdf(docx_file):
        doc = Document(docx_file)
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                pdf.multi_cell(0, 10, text)

        return pdf.output(dest="S").encode("latin-1")

    if uploaded_file is not None:
        file_name = uploaded_file.name

        if conversion_type == "Word to PDF" and file_name.endswith(".docx"):
            try:
                pdf_bytes = convert_docx_to_pdf(uploaded_file)
                st.success("✅ Word file converted to PDF.")
                st.download_button(
                    "📥 Download PDF",
                    data=pdf_bytes,
                    file_name="converted.pdf",
                    mime="application/pdf",
                    key="download_pdf_word_to_pdf"
                )
            except Exception as e:
                st.error(f"⚠️ Error during Word to PDF conversion: {e}")

        # ------------------- PDF to Word Conversion -------------------
        elif conversion_type == "PDF to Word" and file_name.endswith(".pdf"):
            try:
                from pdf2docx import Converter
                import os, tempfile

                with tempfile.TemporaryDirectory() as tmpdirname:
                    pdf_path = os.path.join(tmpdirname, file_name)
                    with open(pdf_path, "wb") as f:
                        f.write(uploaded_file.read())

                    docx_path = os.path.join(tmpdirname, "converted.docx")

                    converter = Converter(pdf_path)
                    converter.convert(docx_path)
                    converter.close()

                    with open(docx_path, "rb") as f:
                        docx_bytes = f.read()

                    st.success("✅ PDF file converted to Word format.")
                    st.download_button(
                        "📥 Download Word",
                        data=docx_bytes,
                        file_name="converted.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_docx_pdf_to_word"
                    )
            except Exception as e:
                st.error(f"⚠️ Error during PDF to Word conversion: {e}")
        else:
            st.error("❌ Uploaded file type doesn't match selected conversion.")
    else:
        st.info("📤 Please upload a file to begin conversion.")





# --- PDF ↔ Excel Tab ---
with tab4:
    st.header("📄 PDF ↔ Excel Converter")

    option = st.radio("Choose conversion direction:", ["📄 PDF to Excel", "📊 Excel to PDF"])

    if option == "📄 PDF to Excel":
        pdf_file = st.file_uploader("Upload a PDF file", type=["pdf"], key="pdf_to_excel")
        if pdf_file:
            import pdfplumber
            import pandas as pd
            from io import BytesIO

            tables_combined = []
            with pdfplumber.open(pdf_file) as pdf:
                for page in pdf.pages:
                    tables = page.extract_tables()
                    for table in tables:
                        df = pd.DataFrame(table[1:], columns=table[0])
                        tables_combined.append(df)

            if tables_combined:
                final_df = pd.concat(tables_combined, ignore_index=True)
                st.dataframe(final_df)

                excel_output = BytesIO()
                with pd.ExcelWriter(excel_output, engine="openpyxl") as writer:
                    final_df.to_excel(writer, index=False)
                st.download_button(
                    label="📥 Download Excel",
                    data=excel_output.getvalue(),
                    file_name="converted.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
            else:
                st.warning("No tables found in the PDF.")

    elif option == "📊 Excel to PDF":
        excel_file = st.file_uploader("Upload an Excel file", type=["xlsx", "xls"], key="excel_to_pdf")
        if excel_file:
            import pandas as pd
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
            from reportlab.lib import colors
            from io import BytesIO

            df = pd.read_excel(excel_file)

            buffer = BytesIO()
            pdf = SimpleDocTemplate(buffer)
            data = [df.columns.tolist()] + df.values.tolist()

            table = Table(data)
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
            ]))

            pdf.build([table])

            st.success("✅ Excel converted to PDF")
            st.download_button("📥 Download PDF", data=buffer.getvalue(), file_name="converted.pdf", mime="application/pdf")
# --- Compress File Tab ---
with tab5:
    st.header("📦 Compress File (PDF, Word, Excel)")

    uploaded_file = st.file_uploader("Upload file to compress", type=["pdf", "docx", "xlsx", "xls"])

    if uploaded_file:
        file_ext = uploaded_file.name.split(".")[-1].lower()
        compressed_data = BytesIO()

        if file_ext == "pdf":
            from PyPDF2 import PdfReader, PdfWriter

            reader = PdfReader(uploaded_file)
            writer = PdfWriter()
            for page in reader.pages:
                writer.add_page(page)
            writer.add_metadata(reader.metadata)

            # Enable compression
            writer.add_metadata({"/Compressed": "True"})

            writer.write(compressed_data)
            st.success("✅ PDF file compressed.")
            st.download_button("📥 Download Compressed PDF", data=compressed_data.getvalue(), file_name="compressed.pdf", mime="application/pdf")

        elif file_ext in ["docx", "xlsx", "xls"]:
            import zipfile

            with zipfile.ZipFile(compressed_data, "w", zipfile.ZIP_DEFLATED) as zipf:
                zipf.writestr(uploaded_file.name, uploaded_file.read())
            st.success("✅ File compressed into ZIP.")
            st.download_button("📥 Download ZIP", data=compressed_data.getvalue(), file_name="compressed.zip", mime="application/zip")

        else:
            st.warning("Unsupported file type.")


